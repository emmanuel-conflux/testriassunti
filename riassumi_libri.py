#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
riassumi_libri.py - CLI Tool per Riassunti di Libri via Ollama (OTTIMIZZATO)

Legge file EPUB o PDF, li analizza tramite Ollama locale e genera
riassunti dettagliati capitolo per capitolo in formato DOCX e Markdown.

Funzionalità:
- Filtro intelligente front/back matter
- Merge di sezioni corte
- Resume robusto con checkpoint
- Modalità dry-run
- Parametri configurabili per ottimizzazione velocità
"""

import os
import sys
import argparse
import time
import re
import json
import hashlib
from pathlib import Path
from typing import List, Dict, Tuple, Optional
from dataclasses import dataclass

# Dependencies
import requests
from tqdm import tqdm

# EPUB handling
try:
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup
    EPUB_SUPPORT = True
except ImportError:
    EPUB_SUPPORT = False
    print("⚠️  ebooklib o BeautifulSoup non disponibili. Supporto EPUB disabilitato.")

# PDF handling
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    print("⚠️  PyPDF2 non disponibile. Supporto PDF disabilitato.")

# DOCX generation
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    print("⚠️  python-docx non disponibile. Generazione DOCX disabilitata.")


# ============================================================================
# CONFIGURAZIONE DEFAULT
# ============================================================================

DEFAULT_INPUT_DIR = r"C:\dariassumere"
DEFAULT_OUTPUT_DIR = r"C:\riassunti"
DEFAULT_MODEL = "qwen3:8b"
DEFAULT_LANGUAGE = "it"

# Filtro capitoli
DEFAULT_MIN_WORDS = 1000           # Soglia minima parole per capitolo valido
DEFAULT_MIN_WORDS_MERGE = 500      # Soglia per merge
DEFAULT_TARGET_WORDS_MERGE = 1500  # Target parole dopo merge

# Chunking
DEFAULT_MAX_CHARS = 16000          # Max caratteri per chunk
DEFAULT_OVERLAP = 400              # Overlap tra chunk

# Verbosità LLM
DEFAULT_MAP_WORDS = 300            # Parole target per MAP
DEFAULT_REDUCE_WORDS = 550         # Parole target per REDUCE
DEFAULT_NUM_PREDICT = 800          # Limitazione token output LLM

OLLAMA_URL = "http://localhost:11434/api/generate"

# Regex per escludere front/back matter (case-insensitive)
SKIP_TITLE_RE = re.compile(
    r'^(cover|front(\s*|-)cover|copyright|toc|table\s+of\s+contents|'
    r'dedication|foreword|preface|acknowledg(e)?ments?|'
    r'notes?|footnotes?|endnotes?|references?|bibliography|'
    r'index|glossary|about(\s+the)?\s+author|about(\s+the)?\s+editors?)$',
    re.I
)


# ============================================================================
# DATACLASSES
# ============================================================================

@dataclass
class ChapterInfo:
    """Informazioni su un capitolo."""
    index: int
    title: str
    text: str
    word_count: int
    hash: str
    status: str = "pending"  # pending, kept, merged, skipped


@dataclass
class ResumeState:
    """Stato di resume per un libro."""
    completed: List[int]
    chapter_hashes: Dict[int, str]
    total_chapters: int
    model: str

    def to_dict(self) -> dict:
        return {
            "completed": self.completed,
            "chapter_hashes": self.chapter_hashes,
            "total_chapters": self.total_chapters,
            "model": self.model
        }

    @staticmethod
    def from_dict(data: dict) -> 'ResumeState':
        return ResumeState(
            completed=data.get("completed", []),
            chapter_hashes={int(k): v for k, v in data.get("chapter_hashes", {}).items()},
            total_chapters=data.get("total_chapters", 0),
            model=data.get("model", DEFAULT_MODEL)
        )


# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def ensure_directory(path: str) -> None:
    """Crea la directory se non esiste."""
    Path(path).mkdir(parents=True, exist_ok=True)


def sanitize_filename(filename: str) -> str:
    """Rimuove caratteri non validi dal nome file."""
    name = Path(filename).stem
    name = re.sub(r'[<>:"/\\|?*]', '_', name)
    return name.strip()


def count_words(text: str) -> int:
    """Conta le parole in un testo."""
    return len(text.split())


def compute_hash(text: str) -> str:
    """Calcola hash MD5 di un testo (8 caratteri)."""
    return hashlib.md5(text.encode('utf-8')).hexdigest()[:8]


def should_skip_section(title: str) -> bool:
    """Verifica se una sezione deve essere saltata (front/back matter)."""
    title_clean = title.strip()
    return SKIP_TITLE_RE.match(title_clean) is not None


def chunk_text(text: str, max_size: int, overlap: int) -> List[str]:
    """
    Suddivide il testo in blocchi di dimensione massima con overlap.

    Args:
        text: Testo da suddividere
        max_size: Dimensione massima di ogni blocco in caratteri
        overlap: Numero di caratteri di sovrapposizione tra blocchi

    Returns:
        Lista di blocchi di testo
    """
    if len(text) <= max_size:
        return [text]

    chunks = []
    start = 0

    while start < len(text):
        end = start + max_size

        # Se non siamo alla fine, cerca un punto di divisione naturale
        if end < len(text):
            search_start = max(start, end - overlap)
            natural_break = max(
                text.rfind('.', search_start, end),
                text.rfind('\n', search_start, end),
                text.rfind(' ', search_start, end)
            )
            if natural_break > start:
                end = natural_break + 1

        chunks.append(text[start:end].strip())
        start = end - overlap if end < len(text) else end

    return chunks


# ============================================================================
# FILTRO E MERGE SEZIONI
# ============================================================================

def filter_and_merge_sections(
    sections: List[Tuple[str, str]],
    min_words: int,
    min_words_merge: int,
    target_words_merge: int,
    dry_run: bool = False
) -> List[ChapterInfo]:
    """
    Filtra front/back matter e unisce sezioni corte.

    Args:
        sections: Lista di tuple (title, text)
        min_words: Soglia minima parole per capitolo valido
        min_words_merge: Soglia per merge
        target_words_merge: Target parole dopo merge
        dry_run: Se True, stampa info senza elaborare

    Returns:
        Lista di ChapterInfo
    """
    chapters = []

    # Step 1: Filtra front/back matter e crea ChapterInfo
    filtered = []
    for idx, (title, text) in enumerate(sections, 1):
        word_count = count_words(text)
        text_hash = compute_hash(text)

        # Verifica se saltare per titolo
        if should_skip_section(title):
            if dry_run:
                chapters.append(ChapterInfo(
                    index=idx,
                    title=title,
                    text=text,
                    word_count=word_count,
                    hash=text_hash,
                    status="skipped (front/back matter)"
                ))
            continue

        # Verifica soglia minima
        if word_count < min_words:
            if dry_run:
                chapters.append(ChapterInfo(
                    index=idx,
                    title=title,
                    text=text,
                    word_count=word_count,
                    hash=text_hash,
                    status="skipped (< min_words)"
                ))
            else:
                # Aggiungi comunque per eventuale merge
                filtered.append((title, text, word_count, text_hash, idx))
            continue

        filtered.append((title, text, word_count, text_hash, idx))

    if dry_run:
        return chapters

    # Step 2: Merge sezioni corte consecutive
    merged = []
    buf_title, buf_texts, buf_words, buf_hashes, buf_indices = None, [], 0, [], []

    def flush():
        nonlocal buf_title, buf_texts, buf_words, buf_hashes, buf_indices
        if buf_title is not None:
            combined_text = "\n\n".join(buf_texts)
            combined_hash = compute_hash(combined_text)
            merged.append((
                buf_title,
                combined_text,
                buf_words,
                combined_hash,
                buf_indices[0],  # Usa primo indice
                len(buf_indices) > 1  # Flag merged
            ))
        buf_title, buf_texts, buf_words, buf_hashes, buf_indices = None, [], 0, [], []

    for title, text, words, text_hash, idx in filtered:
        if buf_title is None:
            buf_title, buf_texts, buf_words, buf_hashes, buf_indices = title, [text], words, [text_hash], [idx]
        else:
            # Merge se non abbiamo ancora raggiunto target o sezione corrente è troppo corta
            if buf_words < target_words_merge or words < min_words_merge:
                buf_texts.append(text)
                buf_words += words
                buf_hashes.append(text_hash)
                buf_indices.append(idx)
            else:
                flush()
                buf_title, buf_texts, buf_words, buf_hashes, buf_indices = title, [text], words, [text_hash], [idx]

    flush()

    # Step 3: Converti in ChapterInfo
    for idx, (title, text, words, text_hash, orig_idx, was_merged) in enumerate(merged, 1):
        status = "kept (merged)" if was_merged else "kept"
        chapters.append(ChapterInfo(
            index=idx,
            title=title,
            text=text,
            word_count=words,
            hash=text_hash,
            status=status
        ))

    return chapters


# ============================================================================
# OLLAMA INTERACTION
# ============================================================================

def call_ollama(
    prompt: str,
    model: str,
    temperature: float = 0.3,
    num_ctx: int = 32768,
    num_predict: int = DEFAULT_NUM_PREDICT,
    timeout: int = 180,
    max_retries: int = 3
) -> Optional[str]:
    """
    Chiama l'API Ollama con retry logic e backoff esponenziale.

    Args:
        prompt: Il prompt da inviare
        model: Nome del modello Ollama
        temperature: Temperatura per la generazione
        num_ctx: Context window
        num_predict: Limitazione token output
        timeout: Timeout in secondi
        max_retries: Numero massimo di tentativi

    Returns:
        Testo generato o None in caso di errore
    """
    payload = {
        "model": model,
        "prompt": prompt,
        "stream": False,
        "options": {
            "temperature": temperature,
            "num_ctx": num_ctx,
            "num_predict": num_predict
        }
    }

    for attempt in range(max_retries):
        try:
            response = requests.post(OLLAMA_URL, json=payload, timeout=timeout)
            response.raise_for_status()

            result = response.json()
            return result.get("response", "").strip()

        except requests.exceptions.RequestException as e:
            if attempt < max_retries - 1:
                wait_time = 2 ** (attempt + 1)  # 2s, 4s, 8s
                print(f"      ⚠️  Tentativo {attempt + 1} fallito: {e}")
                print(f"         Riprovo tra {wait_time} secondi...")
                time.sleep(wait_time)
            else:
                print(f"      ❌ Errore Ollama dopo {max_retries} tentativi: {e}")
                return None
        except Exception as e:
            print(f"      ❌ Errore imprevisto: {e}")
            return None

    return None


# ============================================================================
# EPUB EXTRACTION
# ============================================================================

def extract_sections_from_epub(filepath: str) -> List[Tuple[str, str]]:
    """
    Estrae sezioni da un file EPUB (senza filtri, ritorna tutte le sezioni).

    Args:
        filepath: Percorso del file EPUB

    Returns:
        Lista di tuple (title, text)
    """
    if not EPUB_SUPPORT:
        print("❌ Supporto EPUB non disponibile. Installa: pip install ebooklib beautifulsoup4")
        return []

    try:
        book = epub.read_epub(filepath)
        sections = []

        # Ottieni gli item dello spine (ordine di lettura)
        spine_items = [book.get_item_with_id(item_id) for item_id, _ in book.spine]

        for idx, item in enumerate(spine_items, 1):
            if item is None:
                continue

            try:
                content = item.get_content()
                soup = BeautifulSoup(content, 'html.parser')

                # Rimuovi elementi non testuali
                for tag in soup(['script', 'style', 'svg', 'img']):
                    tag.decompose()

                # Estrai testo
                text = soup.get_text(separator='\n', strip=True)

                if not text.strip():
                    continue

                # Prova a estrarre il titolo
                title = f"Sezione {idx}"
                h1 = soup.find(['h1', 'h2', 'h3'])
                if h1:
                    title = h1.get_text(strip=True) or title

                sections.append((title, text))

            except Exception as e:
                print(f"      ⚠️  Errore nell'elaborazione item {idx}: {e}")
                continue

        return sections

    except Exception as e:
        print(f"❌ Errore nella lettura EPUB: {e}")
        return []


# ============================================================================
# PDF EXTRACTION
# ============================================================================

def extract_sections_from_pdf(filepath: str) -> List[Tuple[str, str]]:
    """
    Estrae sezioni da un file PDF (senza filtri, ritorna tutte le sezioni).

    Args:
        filepath: Percorso del file PDF

    Returns:
        Lista di tuple (title, text)
    """
    if not PDF_SUPPORT:
        print("❌ Supporto PDF non disponibile. Installa: pip install PyPDF2")
        return []

    try:
        with open(filepath, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            full_text = ""

            for page in reader.pages:
                try:
                    full_text += page.extract_text() + "\n"
                except Exception as e:
                    print(f"      ⚠️  Errore nell'estrazione pagina: {e}")
                    continue

        # Tenta di dividere in capitoli usando pattern comuni
        chapter_patterns = [
            r'(?:^|\n)(?:CAPITOLO|Capitolo|CHAPTER|Chapter)\s+(\d+|[IVXLCDM]+)[:\.\s]+(.*?)(?=\n(?:CAPITOLO|Capitolo|CHAPTER|Chapter)|\Z)',
            r'(?:^|\n)(\d+)\s*[\.\)]\s+(.*?)(?=\n\d+\s*[\.\)]|\Z)',
        ]

        sections = []
        found_chapters = False

        for pattern in chapter_patterns:
            matches = list(re.finditer(pattern, full_text, re.MULTILINE | re.DOTALL))
            if matches:
                for idx, match in enumerate(matches):
                    title = match.group(0).split('\n')[0].strip()

                    # Trova il testo del capitolo
                    start_pos = match.start()
                    end_pos = matches[idx + 1].start() if idx + 1 < len(matches) else len(full_text)
                    text = full_text[start_pos:end_pos].strip()

                    if text:
                        sections.append((title if title else f"Sezione {idx + 1}", text))

                if sections:
                    found_chapters = True
                    break

        # Se non trova capitoli, divide in blocchi di dimensione fissa
        if not found_chapters:
            print("      ⚠️  Nessun pattern di capitolo trovato. Suddivisione in blocchi...")
            words = full_text.split()
            chunk_size = 3000  # parole per blocco

            for i in range(0, len(words), chunk_size):
                chunk_words = words[i:i + chunk_size]
                text = ' '.join(chunk_words)

                if text.strip():
                    sections.append((f"Sezione {len(sections) + 1}", text))

        return sections

    except Exception as e:
        print(f"❌ Errore nella lettura PDF: {e}")
        return []


# ============================================================================
# RESUME STATE MANAGEMENT
# ============================================================================

def load_resume_state(state_path: str) -> Optional[ResumeState]:
    """Carica lo stato di resume da file JSON."""
    if not os.path.exists(state_path):
        return None

    try:
        with open(state_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return ResumeState.from_dict(data)
    except Exception as e:
        print(f"      ⚠️  Errore nel caricamento state.json: {e}")
        return None


def save_resume_state(state: ResumeState, state_path: str) -> None:
    """Salva lo stato di resume su file JSON."""
    try:
        with open(state_path, 'w', encoding='utf-8') as f:
            json.dump(state.to_dict(), f, indent=2, ensure_ascii=False)
    except Exception as e:
        print(f"      ⚠️  Errore nel salvataggio state.json: {e}")


def scan_existing_chapters(chapters_dir: str) -> Dict[int, str]:
    """
    Scansiona i capitoli già presenti nella directory.

    Args:
        chapters_dir: Directory dei capitoli

    Returns:
        Dizionario {index: hash} dei capitoli trovati
    """
    existing = {}

    if not os.path.exists(chapters_dir):
        return existing

    # Pattern: NN_<titolo>_<hash8>.md
    pattern = re.compile(r'^(\d+)_.*_([a-f0-9]{8})\.md$')

    for filename in os.listdir(chapters_dir):
        match = pattern.match(filename)
        if match:
            idx = int(match.group(1))
            hash_val = match.group(2)
            existing[idx] = hash_val

    return existing


# ============================================================================
# SUMMARIZATION
# ============================================================================

def summarize_chapter(
    chapter_text: str,
    chapter_title: str,
    model: str,
    max_chars: int,
    overlap: int,
    map_words: int,
    reduce_words: int,
    num_predict: int
) -> Optional[str]:
    """
    Riassume un capitolo usando MAP-REDUCE se necessario.

    Args:
        chapter_text: Testo del capitolo
        chapter_title: Titolo del capitolo
        model: Modello Ollama da usare
        max_chars: Max caratteri per chunk
        overlap: Overlap tra chunk
        map_words: Target parole per MAP
        reduce_words: Target parole per REDUCE
        num_predict: Limitazione token output

    Returns:
        Riassunto del capitolo o None
    """
    # Prompt MAP
    prompt_map = f"""Sei un analista testuale.
Il testo può essere in italiano o inglese, ma rispondi solo in italiano.
Crea un riassunto dettagliato del seguente frammento (IMPORTANTE: scrivi ~{map_words} parole, non superare di molto).

# Sintesi
# Temi chiave
# Personaggi/Concetti principali
# Citazioni (se presenti)
# Osservazioni sullo stile

FRAMMENTO:
{{text}}

RIASSUNTO IN ITALIANO (~{map_words} parole):"""

    # Prompt REDUCE
    prompt_reduce = f"""Unisci e armonizza i seguenti riassunti parziali del capitolo.
Rispondi in italiano, producendo un riassunto coerente (IMPORTANTE: scrivi ~{reduce_words} parole, non superare di molto).

RIASSUNTI PARZIALI:
{{summaries}}

RIASSUNTO UNIFICATO IN ITALIANO (~{reduce_words} parole):"""

    # Se il testo è abbastanza corto, usa direttamente MAP
    if len(chapter_text) <= max_chars:
        prompt = prompt_map.replace("{text}", chapter_text)
        return call_ollama(prompt, model, num_predict=num_predict)

    # Altrimenti usa MAP-REDUCE
    print(f"      📄 Capitolo lungo ({len(chapter_text)} char), applico MAP-REDUCE...")

    # MAP: riassumi ogni chunk
    chunks = chunk_text(chapter_text, max_chars, overlap)
    partial_summaries = []

    for idx, chunk in enumerate(chunks, 1):
        print(f"         Elaboro chunk {idx}/{len(chunks)}...")
        prompt = prompt_map.replace("{text}", chunk)
        summary = call_ollama(prompt, model, num_predict=num_predict)

        if summary:
            partial_summaries.append(summary)
        else:
            print(f"         ⚠️  Chunk {idx} saltato per errore")

    if not partial_summaries:
        print("      ❌ Nessun riassunto parziale generato")
        return None

    # REDUCE: unisci i riassunti parziali
    print(f"      🔄 Unisco {len(partial_summaries)} riassunti parziali...")
    combined = "\n\n---\n\n".join(partial_summaries)
    prompt = prompt_reduce.replace("{summaries}", combined)

    return call_ollama(prompt, model, num_predict=num_predict)


def generate_global_summary(
    chapter_summaries: List[Dict[str, str]],
    model: str,
    num_predict: int
) -> Optional[str]:
    """
    Genera un riassunto complessivo del libro.

    Args:
        chapter_summaries: Lista di dizionari con 'title' e 'summary'
        model: Modello Ollama da usare
        num_predict: Limitazione token output

    Returns:
        Riassunto complessivo o None
    """
    summaries_text = ""
    for item in chapter_summaries:
        summaries_text += f"## {item['title']}\n{item['summary']}\n\n"

    prompt = f"""Genera un riassunto complessivo in italiano basato sui seguenti riassunti dei capitoli:

{summaries_text}

Produce un documento strutturato con:
# Trama complessiva
# Temi e messaggi ricorrenti
# Evoluzione dei personaggi/idee
# Citazioni rappresentative
# Stile e tono
# Sintesi finale

RIASSUNTO COMPLESSIVO IN ITALIANO:"""

    return call_ollama(prompt, model, num_predict=num_predict * 2)  # Più spazio per sintesi globale


# ============================================================================
# OUTPUT GENERATION
# ============================================================================

def write_docx_output(
    book_title: str,
    author: Optional[str],
    chapter_summaries: List[Dict[str, str]],
    global_summary: str,
    output_path: str
) -> bool:
    """
    Genera un file DOCX con i riassunti.

    Args:
        book_title: Titolo del libro
        author: Autore (opzionale)
        chapter_summaries: Lista di riassunti dei capitoli
        global_summary: Riassunto complessivo
        output_path: Percorso del file di output

    Returns:
        True se successo, False altrimenti
    """
    if not DOCX_SUPPORT:
        print("⚠️  Generazione DOCX non disponibile. Installa: pip install python-docx")
        return False

    try:
        doc = Document()

        # Copertina
        title_para = doc.add_heading(f'Riassunto dettagliato', 0)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        subtitle = doc.add_heading(book_title, level=1)
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        if author:
            author_para = doc.add_paragraph(author)
            author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_page_break()

        # Indice
        doc.add_heading('Indice', level=1)
        for idx, item in enumerate(chapter_summaries, 1):
            doc.add_paragraph(f"{idx}. {item['title']}", style='List Number')

        doc.add_page_break()

        # Capitoli
        for item in chapter_summaries:
            doc.add_heading(item['title'], level=1)
            doc.add_paragraph(item['summary'])
            doc.add_paragraph()

        doc.add_page_break()

        # Sintesi complessiva
        doc.add_heading('Sintesi complessiva del libro', level=1)
        doc.add_paragraph(global_summary)

        # Salva
        doc.save(output_path)
        return True

    except Exception as e:
        print(f"❌ Errore nella generazione DOCX: {e}")
        return False


def write_md_output(
    book_title: str,
    author: Optional[str],
    chapter_summaries: List[Dict[str, str]],
    global_summary: str,
    output_path: str
) -> bool:
    """
    Genera un file Markdown con i riassunti.

    Args:
        book_title: Titolo del libro
        author: Autore (opzionale)
        chapter_summaries: Lista di riassunti dei capitoli
        global_summary: Riassunto complessivo
        output_path: Percorso del file di output

    Returns:
        True se successo, False altrimenti
    """
    try:
        content = f"# Riassunto dettagliato — {book_title}\n\n"

        if author:
            content += f"**Autore:** {author}\n\n"

        # Indice
        content += "## Indice\n\n"
        for idx, item in enumerate(chapter_summaries, 1):
            content += f"{idx}. {item['title']}\n"
        content += "\n---\n\n"

        # Capitoli
        for item in chapter_summaries:
            content += f"## {item['title']}\n\n"
            content += f"{item['summary']}\n\n"
            content += "---\n\n"

        # Sintesi complessiva
        content += "## Sintesi complessiva del libro\n\n"
        content += f"{global_summary}\n"

        # Scrivi file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)

        return True

    except Exception as e:
        print(f"❌ Errore nella generazione Markdown: {e}")
        return False


# ============================================================================
# DRY RUN
# ============================================================================

def print_dry_run_report(chapters: List[ChapterInfo], avg_time_per_chapter: float = 2.5):
    """
    Stampa report di dry-run.

    Args:
        chapters: Lista di ChapterInfo
        avg_time_per_chapter: Tempo medio stimato per capitolo (minuti)
    """
    print("\n" + "="*80)
    print("📊 DRY-RUN REPORT")
    print("="*80)

    # Tabella
    print(f"\n{'#':<5} {'Titolo':<40} {'Parole':<10} {'Status':<30}")
    print("-"*80)

    kept_count = 0
    merged_count = 0
    skipped_count = 0

    for ch in chapters:
        status_display = ch.status
        if "kept" in ch.status.lower():
            kept_count += 1
        elif "merged" in ch.status.lower():
            merged_count += 1
        elif "skipped" in ch.status.lower():
            skipped_count += 1

        print(f"{ch.index:<5} {ch.title[:40]:<40} {ch.word_count:<10} {status_display:<30}")

    print("-"*80)

    # Statistiche
    total_kept = sum(1 for ch in chapters if "kept" in ch.status.lower())
    total_time = total_kept * avg_time_per_chapter

    print(f"\n📈 STATISTICHE:")
    print(f"   Sezioni originali: {len(chapters)}")
    print(f"   Capitoli finali (kept): {total_kept}")
    print(f"   Sezioni merged: {merged_count}")
    print(f"   Sezioni skipped: {skipped_count}")
    print(f"   Tempo stimato: ~{total_time:.1f} minuti ({total_time/60:.1f} ore)")
    print("="*80 + "\n")


# ============================================================================
# MAIN PROCESSING
# ============================================================================

def process_book(
    filepath: str,
    output_dir: str,
    args: argparse.Namespace
) -> bool:
    """
    Elabora un singolo libro con resume support.

    Args:
        filepath: Percorso del file del libro
        output_dir: Directory di output
        args: Argomenti CLI

    Returns:
        True se successo, False altrimenti
    """
    filepath = Path(filepath)
    extension = filepath.suffix.lower()
    book_title = sanitize_filename(filepath.name)

    # Crea directory libro
    book_dir = os.path.join(output_dir, book_title)
    chapters_dir = os.path.join(book_dir, "chapters")
    ensure_directory(chapters_dir)

    state_path = os.path.join(book_dir, "state.json")

    print(f"\n{'='*80}")
    print(f"📚 Elaborazione: {filepath.name}")
    print(f"{'='*80}\n")

    # ========================================================================
    # [1/5] Scansione & parsing
    # ========================================================================
    print(f"[1/5] Scansione & parsing")

    if extension == '.epub':
        sections = extract_sections_from_epub(str(filepath))
    elif extension == '.pdf':
        sections = extract_sections_from_pdf(str(filepath))
    else:
        print(f"❌ Formato non supportato: {extension}")
        return False

    if not sections:
        print("❌ Nessuna sezione estratta")
        return False

    print(f"✅ Estratte {len(sections)} sezioni\n")

    # ========================================================================
    # [2/5] Filtri & merge sezioni
    # ========================================================================
    print(f"[2/5] Filtri & merge sezioni")

    chapters = filter_and_merge_sections(
        sections,
        args.min_words,
        args.min_words_merge,
        args.target_words_merge,
        dry_run=args.dry_run
    )

    if args.dry_run:
        # Dry-run: stampa report e termina
        print_dry_run_report(chapters)
        return True

    # Conta statistiche
    kept = [ch for ch in chapters if "kept" in ch.status.lower()]
    merged = sum(1 for ch in chapters if "merged" in ch.status.lower())
    skipped = len(sections) - len(kept)

    print(f"✅ Capitoli finali: {len(kept)} (merged: {merged}, skipped: {skipped})\n")

    if not kept:
        print("❌ Nessun capitolo valido dopo filtri")
        return False

    # ========================================================================
    # Resume: carica stato
    # ========================================================================
    resume_state = load_resume_state(state_path)
    existing_chapters = scan_existing_chapters(chapters_dir)

    if resume_state is None:
        # Nuovo processing o state.json corrotto/assente
        resume_state = ResumeState(
            completed=[],
            chapter_hashes={},
            total_chapters=len(kept),
            model=args.model
        )
        print(f"🆕 Nuovo processing")
    else:
        print(f"♻️  Resume da checkpoint: {len(resume_state.completed)}/{resume_state.total_chapters} completati")

    # Verifica se ci sono cambiamenti
    if resume_state.model != args.model:
        print(f"⚠️  Modello cambiato ({resume_state.model} → {args.model}), rigenero tutto")
        resume_state.completed = []
        resume_state.chapter_hashes = {}

    # ========================================================================
    # [3/5] Riassunto capitoli
    # ========================================================================
    print(f"\n[3/5] Riassunto capitoli ({len(kept)} totali)")
    chapter_summaries = []

    for ch in kept:
        idx = ch.index
        title = ch.title
        text = ch.text
        text_hash = ch.hash

        # Verifica se già completato
        if idx in resume_state.completed:
            # Verifica hash
            if resume_state.chapter_hashes.get(idx) == text_hash:
                # Verifica file esistente
                chapter_file = None
                for f in os.listdir(chapters_dir):
                    if f.startswith(f"{idx:02d}_") and f.endswith(f"_{text_hash}.md"):
                        chapter_file = f
                        break

                if chapter_file:
                    # Capitolo già completato, carica summary
                    print(f"\n   ⏭️  Salto {idx}/{len(kept)} — già riassunto (hash {text_hash})")

                    chapter_path = os.path.join(chapters_dir, chapter_file)
                    try:
                        with open(chapter_path, 'r', encoding='utf-8') as f:
                            summary = f.read()
                        chapter_summaries.append({
                            'title': title,
                            'summary': summary
                        })
                        continue
                    except Exception as e:
                        print(f"      ⚠️  Errore nel caricamento, rigenero: {e}")
            else:
                print(f"\n   🔄 Capitolo {idx}/{len(kept)} — hash cambiato, rigenero")

        # Genera riassunto
        print(f"\n   📘 Capitolo {idx}/{len(kept)} — {title} (~{ch.word_count} parole)")

        summary = summarize_chapter(
            text,
            title,
            args.model,
            args.max_chars,
            args.overlap,
            args.map_words,
            args.reduce_words,
            args.num_predict
        )

        if summary:
            chapter_summaries.append({
                'title': title,
                'summary': summary
            })

            # Salva capitolo
            safe_title = sanitize_filename(title)
            chapter_filename = f"{idx:02d}_{safe_title}_{text_hash}.md"
            chapter_path = os.path.join(chapters_dir, chapter_filename)

            try:
                with open(chapter_path, 'w', encoding='utf-8') as f:
                    f.write(summary)

                # Aggiorna stato
                if idx not in resume_state.completed:
                    resume_state.completed.append(idx)
                resume_state.chapter_hashes[idx] = text_hash
                save_resume_state(resume_state, state_path)

                print(f"   ✅ Completato")
            except Exception as e:
                print(f"   ⚠️  Errore nel salvataggio: {e}")
        else:
            print(f"   ⚠️  Saltato per errore")

    if not chapter_summaries:
        print("\n❌ Nessun riassunto generato")
        return False

    # ========================================================================
    # [4/5] Sintesi globale
    # ========================================================================
    print(f"\n[4/5] Sintesi globale")
    global_summary = generate_global_summary(chapter_summaries, args.model, args.num_predict)

    if not global_summary:
        print("⚠️  Riassunto complessivo non generato, uso sintesi base")
        global_summary = "Riassunto complessivo non disponibile."
    else:
        print("✅ Riassunto complessivo generato")

    # ========================================================================
    # [5/5] Scrittura output
    # ========================================================================
    print(f"\n[5/5] Scrittura output")

    docx_path = os.path.join(book_dir, "final.docx")
    md_path = os.path.join(book_dir, "final.md")

    docx_ok = write_docx_output(book_title, None, chapter_summaries, global_summary, docx_path)
    md_ok = write_md_output(book_title, None, chapter_summaries, global_summary, md_path)

    if docx_ok:
        print(f"✅ DOCX: {docx_path}")
    if md_ok:
        print(f"✅ MD: {md_path}")

    return docx_ok or md_ok


# ============================================================================
# MAIN CLI
# ============================================================================

def main():
    """Funzione principale del programma."""
    parser = argparse.ArgumentParser(
        description="CLI Tool per Riassunti di Libri via Ollama (OTTIMIZZATO)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Esempi:
  python riassumi_libri.py
  python riassumi_libri.py --model qwen3:30b
  python riassumi_libri.py --dry_run
  python riassumi_libri.py --min_words 800 --map_words 250
        """
    )

    # Parametri principali
    parser.add_argument('--model', type=str, default=DEFAULT_MODEL,
                       help=f'Modello Ollama (default: {DEFAULT_MODEL})')
    parser.add_argument('--input_dir', type=str, default=DEFAULT_INPUT_DIR,
                       help=f'Directory input (default: {DEFAULT_INPUT_DIR})')
    parser.add_argument('--output_dir', type=str, default=DEFAULT_OUTPUT_DIR,
                       help=f'Directory output (default: {DEFAULT_OUTPUT_DIR})')
    parser.add_argument('--language', type=str, default=DEFAULT_LANGUAGE,
                       help=f'Lingua output (default: {DEFAULT_LANGUAGE})')

    # Filtro capitoli
    parser.add_argument('--min_words', type=int, default=DEFAULT_MIN_WORDS,
                       help=f'Parole minime per capitolo valido (default: {DEFAULT_MIN_WORDS})')
    parser.add_argument('--min_words_merge', type=int, default=DEFAULT_MIN_WORDS_MERGE,
                       help=f'Soglia parole per merge (default: {DEFAULT_MIN_WORDS_MERGE})')
    parser.add_argument('--target_words_merge', type=int, default=DEFAULT_TARGET_WORDS_MERGE,
                       help=f'Target parole dopo merge (default: {DEFAULT_TARGET_WORDS_MERGE})')

    # Chunking
    parser.add_argument('--max_chars', type=int, default=DEFAULT_MAX_CHARS,
                       help=f'Max caratteri per chunk (default: {DEFAULT_MAX_CHARS})')
    parser.add_argument('--overlap', type=int, default=DEFAULT_OVERLAP,
                       help=f'Overlap tra chunk (default: {DEFAULT_OVERLAP})')

    # Verbosità LLM
    parser.add_argument('--map_words', type=int, default=DEFAULT_MAP_WORDS,
                       help=f'Parole target per MAP (default: {DEFAULT_MAP_WORDS})')
    parser.add_argument('--reduce_words', type=int, default=DEFAULT_REDUCE_WORDS,
                       help=f'Parole target per REDUCE (default: {DEFAULT_REDUCE_WORDS})')
    parser.add_argument('--num_predict', type=int, default=DEFAULT_NUM_PREDICT,
                       help=f'Limitazione token output LLM (default: {DEFAULT_NUM_PREDICT})')

    # Dry-run
    parser.add_argument('--dry_run', action='store_true',
                       help='Modalità dry-run (mostra solo report senza elaborare)')

    args = parser.parse_args()

    print("\n" + "="*80)
    print("📚 RIASSUMI LIBRI - CLI Tool via Ollama (OTTIMIZZATO)")
    print("="*80)
    print(f"Modello: {args.model}")
    print(f"Input: {args.input_dir}")
    print(f"Output: {args.output_dir}")
    print(f"Min parole/capitolo: {args.min_words}")
    print(f"Min parole merge: {args.min_words_merge}")
    print(f"Target parole merge: {args.target_words_merge}")
    print(f"Max chars/chunk: {args.max_chars}")
    print(f"Overlap: {args.overlap}")
    print(f"MAP words: {args.map_words}")
    print(f"REDUCE words: {args.reduce_words}")
    print(f"Num predict: {args.num_predict}")
    print(f"Dry-run: {args.dry_run}")
    print("="*80 + "\n")

    # Verifica Ollama (salta se dry-run)
    if not args.dry_run:
        print("🔍 Verifica connessione a Ollama...")
        try:
            response = requests.get("http://localhost:11434/api/tags", timeout=5)
            response.raise_for_status()
            print("✅ Ollama raggiungibile\n")
        except Exception as e:
            print(f"❌ Errore connessione Ollama: {e}")
            print("   Assicurati che Ollama sia in esecuzione su http://localhost:11434")
            sys.exit(1)

    # Verifica directory input
    if not os.path.exists(args.input_dir):
        print(f"❌ Directory input non trovata: {args.input_dir}")
        sys.exit(1)

    # Crea directory output
    ensure_directory(args.output_dir)

    # Scansiona file
    print(f"📁 Scansione {args.input_dir}")
    files = []

    for ext in ['.epub', '.pdf']:
        files.extend(Path(args.input_dir).glob(f'*{ext}'))

    if not files:
        print("❌ Nessun file EPUB o PDF trovato")
        sys.exit(1)

    print(f"Trovati {len(files)} file: {', '.join([f.name for f in files])}\n")

    # Elabora ogni file
    success_count = 0

    for idx, filepath in enumerate(files, 1):
        print(f"\n{'#'*80}")
        print(f"FILE {idx}/{len(files)}")
        print(f"{'#'*80}")

        try:
            if process_book(str(filepath), args.output_dir, args):
                success_count += 1
        except KeyboardInterrupt:
            print("\n\n⚠️  Interruzione utente. Progress salvato.")
            sys.exit(0)
        except Exception as e:
            print(f"\n❌ Errore nell'elaborazione: {e}")
            import traceback
            traceback.print_exc()
            continue

    # Riepilogo finale
    if not args.dry_run:
        print(f"\n{'='*80}")
        print(f"✅ OPERAZIONE COMPLETATA")
        print(f"{'='*80}")
        print(f"File elaborati: {success_count}/{len(files)}")
        print(f"Output salvati in: {args.output_dir}")
        print("="*80 + "\n")


if __name__ == "__main__":
    main()
